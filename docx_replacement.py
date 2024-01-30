# Import libraries
from docx import Document
import tkinter as tk
import ttkbootstrap as ttk
from tkinter import filedialog

# Define a function to change text while preserving formatting
def set_text_preserving_text_formatting(paragraph, text):
    runs = paragraph.runs
    runs[0].text = text

# Define a function to replace text in tables
def replace_table_text(doc, old_word, new_word):
    # Get the tables in the document
    tables = doc.tables
    # Loop through each table
    for table in tables:
        # Loop through each row in the table
        for row in table.rows:
            # Loop through each cell in the row
            for cell in row.cells:
                # Loop through each paragraph in the cell
                for paragraph in cell.paragraphs:
                    # If the old word is in the paragraph text, replace it with the new word
                    if old_word in paragraph.text:
                        set_text_preserving_text_formatting(paragraph, paragraph.text.replace(old_word, new_word))

# Define a function to select a file
def browse():
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename()
    file_path = file_path.replace('/', '\\\\')
    filepath.set(file_path)

# Define a function to replace text
def replace():
    old=old_word.get()
    new=new_word.get()
    # Open the document
    doc = Document(filepath.get())
    # Loop through each paragraph in the document
    for p in doc.paragraphs:
        # If the old word is in the paragraph text, replace it with the new word
        if old in p.text:
            set_text_preserving_text_formatting(p, p.text.replace(old, new))
    # Replace text in tables
    replace_table_text(doc,old,new)        
    # Save changes
    doc.save(filepath.get())

# Create a window
window=ttk.Window(themename='darkly')
window.title('word replacement')
window.geometry('400x300')

# Add a header label
top_label=ttk.Label(master=window,text='Browse the file you want '+"\n"+' type the old word and the new word '+"\n"+' then click replace button',font='Calibri 12 bold')
top_label.pack(side='top')

# Add a frame for the file path selection
filepath_form=ttk.Frame(master=window)
filepath_form.pack()
filepath=tk.StringVar()
path_entry=ttk.Entry(master=filepath_form,textvariable=filepath)
path_entry.pack(side='left')
browse_button=ttk.Button(master=filepath_form,text='Browse',command=browse)
browse_button.pack(side='left')

# Add a frame for the old word entry
oldword_form=ttk.Frame(master=window)
oldword_form.pack()
oldword_label=ttk.Label(master=oldword_form,text='old_word:',font='Calibri 10')
oldword_label.pack(side='left',padx=10)
old_word=tk.StringVar()
oldword_entry=ttk.Entry(master=oldword_form,textvariable=old_word)
oldword_entry.pack(side='left',pady=10)

# Add a frame for the new word entry
newword_form=ttk.Frame(master=window)
newword_form.pack()
newword_label=ttk.Label(master=newword_form,text='new_word:',font='Calibri 10')
newword_label.pack(side='left',padx=10)
new_word=tk.StringVar()
newword_entry=ttk.Entry(master=newword_form,textvariable=new_word)
newword_entry.pack(side='left',pady=10)

# Add a replace button
replace_button=ttk.Button(master=window,text='Replace',command=replace)
replace_button.pack(pady=10)

# Run the app
window.mainloop()
